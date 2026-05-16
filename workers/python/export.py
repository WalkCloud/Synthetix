"""Synthetix document export — converts Markdown to PDF (HTML-based) or DOCX.

Usage:
  python export.py --input <file.md> --output <output_path> --format [pdf|docx]

For PDF: generates a print-friendly standalone HTML file (browser prints to PDF).
For DOCX: uses python-docx for native Word format with CJK professional typesetting.
"""
import sys
import json
import argparse
import os


def export_to_html(md_path: str, output_path: str) -> dict:
    try:
        import markdown
    except ImportError:
        return {"error": "markdown package not installed. Run: pip install markdown"}

    with open(md_path, "r", encoding="utf-8") as f:
        md_content = f.read()

    title = ""
    for line in md_content.split("\n"):
        if line.startswith("# "):
            title = line[2:].strip()
            break

    html_body = markdown.markdown(
        md_content,
        extensions=["tables", "fenced_code", "codehilite", "toc", "nl2br"],
    )

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<title>{title or "Document"}</title>
<style>
  @media print {{
    body {{ margin: 20mm 25mm; }}
    @page {{ size: A4; }}
  }}
  body {{
    font-family: "Noto Serif CJK SC", "Source Han Serif SC", "SimSun", Georgia, serif;
    font-size: 12pt;
    line-height: 1.8;
    color: #1a1a1a;
    max-width: 210mm;
    margin: 40px auto;
    padding: 0 20px;
  }}
  h1 {{ font-size: 22pt; border-bottom: 2px solid #333; padding-bottom: 8px; margin-top: 0; }}
  h2 {{ font-size: 16pt; margin-top: 28px; border-bottom: 1px solid #ccc; padding-bottom: 4px; }}
  h3 {{ font-size: 14pt; margin-top: 22px; }}
  p {{ margin: 8px 0; }}
  pre {{ background: #f5f5f5; padding: 12px 16px; border-radius: 6px; font-size: 10pt; overflow-x: auto; }}
  code {{ font-family: "JetBrains Mono", "Courier New", monospace; font-size: 10pt; }}
  table {{ border-collapse: collapse; margin: 12px 0; width: 100%; }}
  th, td {{ border: 1px solid #ddd; padding: 6px 10px; text-align: left; font-size: 11pt; }}
  th {{ background: #f0f0f0; font-weight: 600; }}
  blockquote {{ border-left: 3px solid #7c3aed; margin: 12px 0; padding: 4px 16px; color: #555; background: #fafafa; }}
  img {{ max-width: 100%; height: auto; }}
  .page-break {{ page-break-before: always; }}
</style>
</head>
<body>
{html_body}
</body>
</html>"""

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)

    return {"status": "ok", "format": "pdf_html", "output": output_path, "title": title}


# ---------------------------------------------------------------------------
# OOXML helpers for CJK-professional DOCX output
# ---------------------------------------------------------------------------

_EA_BODY = "\u5b8b\u4f53"
_EA_HEADING = "\u9ed1\u4f53"
_LATIN_BODY = "Times New Roman"
_LATIN_HEADING = "Times New Roman"
_THEME_ATTRS = ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme")


def _ns(tag: str) -> str:
    return f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{tag}"


def _setup_styles(doc):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    sizes = {
        "Title": (22, _EA_HEADING, _LATIN_HEADING),
        "Heading 1": (16, _EA_HEADING, _LATIN_HEADING),
        "Heading 2": (15, _EA_HEADING, _LATIN_HEADING),
        "Heading 3": (14, _EA_HEADING, _LATIN_HEADING),
        "Heading 4": (13, _EA_BODY, _LATIN_BODY),
        "Heading 5": (12, _EA_BODY, _LATIN_BODY),
        "Normal": (12, _EA_BODY, _LATIN_BODY),
    }

    for style_name, (pt, ea, latin) in sizes.items():
        style = doc.styles[style_name]
        style.font.size = Pt(pt)
        style.font.bold = False
        pf = style.element.find(_ns("rPr"))
        if pf is None:
            from docx.oxml import OxmlElement
            pf = OxmlElement("w:rPr")
            style.element.append(pf)
        _apply_font_element(pf, ea, latin)

    if "Title" in [s.name for s in doc.styles]:
        title_style = doc.styles["Title"]
        title_style.font.bold = True

    normal_style = doc.styles["Normal"]
    normal_style.paragraph_format.line_spacing = 1.5


def _apply_font_element(r_pr, ea: str, latin: str):
    from docx.oxml import OxmlElement

    rFonts = r_pr.find(_ns("rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        r_pr.insert(0, rFonts)

    rFonts.set(_ns("ascii"), latin)
    rFonts.set(_ns("hAnsi"), latin)
    rFonts.set(_ns("eastAsia"), ea)
    rFonts.set(_ns("cs"), latin)

    for attr in _THEME_ATTRS:
        full = _ns(attr)
        if full in rFonts.attrib:
            del rFonts.attrib[full]


def _clean_doc_defaults(doc):
    from docx.oxml import OxmlElement

    styles_element = doc.styles.element
    docDefaults = styles_element.find(_ns("docDefaults"))
    if docDefaults is None:
        return

    rPrDefault = docDefaults.find(_ns("rPrDefault"))
    if rPrDefault is None:
        return
    rPr = rPrDefault.find(_ns("rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rPrDefault.append(rPr)

    _apply_font_element(rPr, _EA_BODY, _LATIN_BODY)

    rFonts = rPr.find(_ns("rFonts"))
    if rFonts is not None:
        for attr in _THEME_ATTRS:
            full = _ns(attr)
            if full in rFonts.attrib:
                del rFonts.attrib[full]


def _setup_heading_numbering(doc):
    from docx.oxml import OxmlElement
    from lxml import etree

    numbering = doc.part.numbering_part.element if hasattr(doc.part, "numbering_part") and doc.part.numbering_part else None
    if numbering is None:
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            numbering_part = doc.part.numbering_part
            numbering = numbering_part.element
        except Exception:
            return

    abstractNums = numbering.findall(_ns("abstractNum"))
    max_abstract_id = max(
        (int(an.get(_ns("abstractNumId"), "0")) for an in abstractNums), default=-1
    )
    new_abstract_id = max_abstract_id + 1

    abstractNum = OxmlElement("w:abstractNum")
    abstractNum.set(_ns("abstractNumId"), str(new_abstract_id))

    lvl_offset = 360
    lvl_indents = [0, 420, 840, 1260, 1680]
    lvl_formats = ["%1.", "%1.%2.", "%1.%2.%3.", "%1.%2.%3.%4.", "%1.%2.%3.%4.%5."]
    heading_styles = ["Heading 1", "Heading 2", "Heading 3", "Heading 4", "Heading 5"]

    for i in range(5):
        lvl = OxmlElement("w:lvl")
        lvl.set(_ns("ilvl"), str(i))
        lvl.set(_ns("tplc"), "04090001")

        startEl = OxmlElement("w:start")
        startEl.set(_ns("val"), "1")
        lvl.append(startEl)

        fmtEl = OxmlElement("w:numFmt")
        fmtEl.set(_ns("val"), "decimal")
        lvl.append(fmtEl)

        txtEl = OxmlElement("w:lvlText")
        txtEl.set(_ns("val"), lvl_formats[i])
        lvl.append(txtEl)

        jmpEl = OxmlElement("w:lvlJc")
        jmpEl.set(_ns("val"), "left")
        lvl.append(jmpEl)

        pPr = OxmlElement("w:pPr")

        indEl = OxmlElement("w:ind")
        indEl.set(_ns("left"), str(lvl_indents[i]))
        indEl.set(_ns("firstLine"), "0")
        pPr.append(indEl)

        numEl = OxmlElement("w:numFmt")
        pPr.append(numEl)

        lvl.append(pPr)

        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(_ns("hint"), "default")
        rPr.append(rFonts)
        lvl.append(rPr)

        abstractNum.append(lvl)

    existing_an = numbering.find(_ns("abstractNum"))
    if existing_an is not None:
        existing_an.addprevious(abstractNum)
    else:
        numbering.append(abstractNum)

    nums = numbering.findall(_ns("num"))
    max_num_id = max((int(n.get(_ns("numId"), "0")) for n in nums), default=0)
    new_num_id = max_num_id + 1

    num = OxmlElement("w:num")
    num.set(_ns("numId"), str(new_num_id))
    abstractRef = OxmlElement("w:abstractNumId")
    abstractRef.set(_ns("val"), str(new_abstract_id))
    num.append(abstractRef)
    numbering.append(num)

    for i, style_name in enumerate(heading_styles):
        style = doc.styles[style_name]
        pPr = style.element.find(_ns("pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            style.element.insert(0, pPr)

        numPr = pPr.find(_ns("numPr"))
        if numPr is not None:
            pPr.remove(numPr)
        numPr = OxmlElement("w:numPr")

        ilvl = OxmlElement("w:ilvl")
        ilvl.set(_ns("val"), str(i))
        numPr.append(ilvl)

        numIdEl = OxmlElement("w:numId")
        numIdEl.set(_ns("val"), str(new_num_id))
        numPr.append(numIdEl)

        pPr.insert(0, numPr)


def _apply_run_font(run, ea: str = _EA_BODY, latin: str = _LATIN_BODY):
    r_elem = run._element
    rPr = r_elem.find(_ns("rPr"))
    if rPr is None:
        from docx.oxml import OxmlElement
        rPr = OxmlElement("w:rPr")
        r_elem.insert(0, rPr)
    _apply_font_element(rPr, ea, latin)


def _add_picture_cn(doc, image_path: str, alt_text: str = ""):
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    try:
        doc.add_picture(image_path, width=Cm(14))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if alt_text:
            from docx.shared import Pt, RGBColor
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_run = cap.add_run(alt_text)
            cap_run.font.size = Pt(9)
            cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            _apply_run_font(cap_run)
    except Exception:
        _add_image_placeholder(doc, alt_text or image_path)


def _add_image_placeholder(doc, caption: str):
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[\u56fe\u7247\u5360\u4f4d\u7b26 \u2014 {caption}\uff1a\u8bf7\u624b\u52a8\u63d2\u5165\u56fe\u7247]")
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    _apply_run_font(run)


def add_md_table(doc, rows: list[str]):
    if len(rows) < 2:
        return

    data = []
    for row in rows:
        cells = [c.strip() for c in row.strip("|").split("|")]
        data.append(cells)

    if not data:
        return

    if all(c.replace("-", "").replace(":", "").strip() == "" for c in data[1] if len(data) > 1):
        header = data[0]
        body = data[2:]
    else:
        header = data[0]
        body = data[1:]

    num_cols = len(header)
    table = doc.add_table(rows=1 + len(body), cols=num_cols)
    table.style = "Light Grid Accent 1"

    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.font.bold = True
        _apply_run_font(run)

    for r, row_data in enumerate(body):
        for c, cell_text in enumerate(row_data):
            if c < num_cols:
                cell = table.rows[r + 1].cells[c]
                run = cell.paragraphs[0].add_run(cell_text)
                _apply_run_font(run)


def _write_inline_formatting(paragraph, text: str):
    import re
    from docx.shared import RGBColor

    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
            _apply_run_font(run)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
            _apply_run_font(run)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            _apply_run_font(run)
        elif part:
            run = paragraph.add_run(part)
            _apply_run_font(run)


def _write_markdown(doc, md_content: str, base_dir: str = "."):
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import base64
    import tempfile

    lines = md_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if line.startswith("# ") and not line.startswith("## "):
            p = doc.add_heading(line[2:], level=1)
            for run in p.runs:
                _apply_run_font(run, _EA_HEADING, _LATIN_HEADING)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
            for run in p.runs:
                _apply_run_font(run, _EA_HEADING, _LATIN_HEADING)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
            for run in p.runs:
                _apply_run_font(run, _EA_HEADING, _LATIN_HEADING)
        elif line.startswith("#### "):
            p = doc.add_heading(line[5:], level=4)
            for run in p.runs:
                _apply_run_font(run, _EA_BODY, _LATIN_BODY)
        elif line.startswith("##### "):
            p = doc.add_heading(line[6:], level=5)
            for run in p.runs:
                _apply_run_font(run, _EA_BODY, _LATIN_BODY)
        elif line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                _apply_run_font(run)
                p.paragraph_format.left_indent = Inches(0.3)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _write_inline_formatting(p, line[2:])
        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            p = doc.add_paragraph(style="List Number")
            _write_inline_formatting(p, text)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            _apply_run_font(run)
        elif line.startswith("---"):
            hr = doc.add_paragraph()
            hr_run = hr.add_run("\u2500" * 60)
            hr_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            _apply_run_font(hr_run)
        elif line == "":
            pass
        elif line.startswith("!["):
            caption_end = line.find("](")
            url_end = line.rfind(")")
            if caption_end > 0 and url_end > caption_end:
                alt_text = line[2:caption_end]
                image_ref = line[caption_end + 2:url_end]

                image_path = None
                if image_ref.startswith("data:image/svg+xml;base64,"):
                    b64_data = image_ref[len("data:image/svg+xml;base64,"):]
                    try:
                        svg_bytes = base64.b64decode(b64_data)
                        tmp = tempfile.NamedTemporaryFile(suffix=".svg", delete=False, mode="wb")
                        tmp.write(svg_bytes)
                        tmp.close()
                        image_path = tmp.name
                    except Exception:
                        pass
                elif image_ref.startswith("data:image/png;base64,"):
                    b64_data = image_ref[len("data:image/png;base64,"):]
                    try:
                        png_bytes = base64.b64decode(b64_data)
                        tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False, mode="wb")
                        tmp.write(png_bytes)
                        tmp.close()
                        image_path = tmp.name
                    except Exception:
                        pass
                else:
                    resolved = os.path.join(base_dir, image_ref)
                    if os.path.exists(resolved):
                        image_path = resolved
                    elif os.path.exists(image_ref):
                        image_path = os.path.abspath(image_ref)

                if image_path and os.path.exists(image_path):
                    _add_picture_cn(doc, image_path, alt_text)
                    if image_ref.startswith("data:"):
                        try:
                            os.unlink(image_path)
                        except OSError:
                            pass
                else:
                    _add_image_placeholder(doc, alt_text)
        elif line.startswith("|"):
            rows = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            if rows:
                add_md_table(doc, rows)
                continue
        else:
            p = doc.add_paragraph()
            _write_inline_formatting(p, line)

        i += 1


def _create_document():
    from docx import Document
    from docx.shared import Cm

    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.4)

    _clean_doc_defaults(doc)
    _setup_styles(doc)

    try:
        _setup_heading_numbering(doc)
    except Exception:
        pass

    return doc


def export_to_docx(md_path: str, output_path: str) -> dict:
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    with open(md_path, "r", encoding="utf-8") as f:
        md_content = f.read()

    doc = _create_document()
    _write_markdown(doc, md_content, base_dir=os.path.dirname(os.path.abspath(md_path)))

    doc.save(output_path)
    return {"status": "ok", "format": "docx", "output": output_path}


def main() -> None:
    parser = argparse.ArgumentParser(description="Synthetix document exporter")
    parser.add_argument("--input", required=True, help="Path to Markdown file")
    parser.add_argument("--output", required=True, help="Output file path")
    parser.add_argument("--format", choices=["pdf", "docx"], required=True)
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(json.dumps({"error": f"Input file not found: {args.input}"}))
        sys.exit(1)

    if args.format == "pdf":
        result = export_to_html(args.input, args.output)
    else:
        result = export_to_docx(args.input, args.output)

    print(json.dumps(result))


if __name__ == "__main__":
    import re
    from docx.shared import RGBColor
    main()
